"""
Testes unitários para gerador/extractors/docx.py
Cria documentos Word em memória com python-docx para testar a extração real.
"""

import io

import pytest
from docx import Document
from docx.shared import Pt

from gerador.extractors.base import TipoArquivo
from gerador.extractors.docx import extrair_docx


# ── Fixtures ──────────────────────────────────────────────────────────────────

def _make_docx(paragrafos: list[str], estilo: str = "Normal") -> io.BytesIO:
    """Cria DOCX em memória com lista de parágrafos."""
    doc = Document()
    for texto in paragrafos:
        doc.add_paragraph(texto, style=estilo)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "documento.docx"
    return buf


def _make_docx_com_questoes() -> io.BytesIO:
    """Cria DOCX simulando uma atividade SEED-PR com questões de múltipla escolha."""
    doc = Document()
    doc.add_heading("Atividade — Front-End", level=1)
    doc.add_paragraph("1. O que é HTML?")
    doc.add_paragraph("A) Uma linguagem de programação")
    doc.add_paragraph("B) Uma linguagem de marcação")
    doc.add_paragraph("C) Um banco de dados")
    doc.add_paragraph("D) Um sistema operacional")
    doc.add_paragraph("2. Qual tag cria um parágrafo em HTML?")
    doc.add_paragraph("A) <div>")
    doc.add_paragraph("B) <span>")
    doc.add_paragraph("C) <p>")
    doc.add_paragraph("D) <br>")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "AULA_03_ATIVIDADE_FRONT_END.docx"
    return buf


def _make_docx_com_tabela() -> io.BytesIO:
    """Cria DOCX com tabela."""
    doc = Document()
    tabela = doc.add_table(rows=2, cols=3)
    tabela.cell(0, 0).text = "Tag"
    tabela.cell(0, 1).text = "Uso"
    tabela.cell(0, 2).text = "Exemplo"
    tabela.cell(1, 0).text = "<p>"
    tabela.cell(1, 1).text = "Parágrafo"
    tabela.cell(1, 2).text = "<p>Texto</p>"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "tabela.docx"
    return buf


# ── Testes ────────────────────────────────────────────────────────────────────

class TestExtrairDocx:

    def test_extrai_paragrafos_simples(self):
        """Deve extrair todos os parágrafos do documento."""
        docx = _make_docx(["Primeiro parágrafo.", "Segundo parágrafo.", "Terceiro."])
        resultado = extrair_docx(docx)

        assert resultado.ok
        assert resultado.tipo == TipoArquivo.DOCX
        assert "Primeiro parágrafo." in resultado.conteudo
        assert "Segundo parágrafo." in resultado.conteudo

    def test_retorna_vazio_para_none(self):
        """Quando arquivo é None, deve retornar ResultadoExtracao sem conteúdo e sem erro."""
        resultado = extrair_docx(None)

        assert resultado.conteudo == ""
        assert resultado.tipo == TipoArquivo.DOCX

    def test_detecta_questoes_numeradas(self):
        """Questões numeradas (1., 1), Questão 1) devem ser detectadas."""
        docx = _make_docx_com_questoes()
        resultado = extrair_docx(docx)

        assert resultado.ok
        assert resultado.metadados.get("questoes_detectadas", 0) >= 2

    def test_detecta_alternativas_multipla_escolha(self):
        """Alternativas A), B), C), D) devem aparecer no conteúdo."""
        docx = _make_docx_com_questoes()
        resultado = extrair_docx(docx)

        assert "A)" in resultado.conteudo or "A." in resultado.conteudo

    def test_extrai_tabelas(self):
        """Tabelas devem ser convertidas para texto com separadores '|'."""
        docx = _make_docx_com_tabela()
        resultado = extrair_docx(docx)

        assert resultado.ok
        assert "Tag" in resultado.conteudo
        assert "|" in resultado.conteudo

    def test_retorna_erro_para_arquivo_invalido(self):
        """Arquivo corrompido deve retornar ResultadoExtracao com erro."""
        invalido = io.BytesIO(b"nao e um docx valido")
        invalido.name = "invalido.docx"
        resultado = extrair_docx(invalido)

        assert not resultado.ok
        assert resultado.erro != ""

    def test_atividade_seed_contem_titulo(self):
        """Documento de atividade deve conter o título extraído."""
        docx = _make_docx_com_questoes()
        resultado = extrair_docx(docx)

        assert "Atividade" in resultado.conteudo or "Front-End" in resultado.conteudo
